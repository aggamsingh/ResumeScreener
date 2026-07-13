"""
CV parsing: PDF and DOCX extraction, OCR fallback, text chunking,
name heuristics, and metadata extraction.

Supported formats: .pdf, .docx
  - PDF: pdfplumber first; falls back to Tesseract OCR if text layer
    is empty or too short (configurable via ENABLE_OCR env var).
  - DOCX: python-docx paragraph extraction.
  - Other: skipped with a log warning (never raises).

Chunking strategy: fixed word-count windows with overlap.
Chosen over section-based parsing for robustness across 20k
heterogeneous CVs from different templates and tools.
"""
from __future__ import annotations

import logging
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from indexer.metadata import CVMetadata, extract_metadata

logger = logging.getLogger(__name__)

# ── Configuration ──────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx"})

CHUNK_WORD_SIZE = 400   # words per chunk
CHUNK_OVERLAP   = 50    # word overlap between adjacent chunks

# Minimum number of characters for extracted text to be considered valid.
# PDFs with fewer chars are treated as scanned and OCR is attempted.
_MIN_TEXT_LENGTH = 100

# Respect the ENABLE_OCR environment variable (default: true if packages available)
_OCR_ENV = os.getenv("ENABLE_OCR", "true").lower()
_OCR_REQUESTED = _OCR_ENV not in ("false", "0", "no", "off")

# Try to import OCR packages; availability determines whether OCR runs
try:
    import pytesseract          # type: ignore
    from pdf2image import convert_from_path  # type: ignore
    from PIL import Image       # type: ignore
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False

OCR_ENABLED = _OCR_REQUESTED and _OCR_AVAILABLE

if _OCR_REQUESTED and not _OCR_AVAILABLE:
    logger.warning(
        "ENABLE_OCR=true but pytesseract/pdf2image/Pillow are not installed. "
        "Scanned PDFs will be skipped. Install OCR dependencies or set ENABLE_OCR=false."
    )

# Keywords that indicate a line is a section header, not a person's name
_HEADER_KEYWORDS = frozenset({
    "resume", "cv", "curriculum", "vitae", "profile", "summary",
    "objective", "contact", "address", "education", "experience",
    "skills", "projects", "references", "declaration",
})


# ── Data Structures ────────────────────────────────────────────────────────────

@dataclass
class ParsedCV:
    """Result of parsing a single CV file, including metadata."""
    candidate_id: str
    name:         str
    cv_path:      str
    raw_text:     str
    metadata:     CVMetadata
    chunks:       list[str] = field(default_factory=list)
    ocr_used:     bool = False   # True if text was obtained via OCR


# ── Name Extraction ────────────────────────────────────────────────────────────

def _extract_name(text: str, filename_stem: str) -> str:
    """
    Heuristic: look for a 2–4 word title-cased line in the first 15 lines.
    Falls back to a cleaned-up version of the filename stem.
    """
    for line in text.splitlines()[:15]:
        line = line.strip()
        if not line:
            continue
        words = line.split()
        if 2 <= len(words) <= 4:
            if all(w[0].isupper() for w in words if w and w[0].isalpha()):
                if not any(kw in line.lower() for kw in _HEADER_KEYWORDS):
                    if not any(ch in line for ch in (":", "|", "/", "@", "–", "!", "?")):
                        return line

    # Fallback: derive from filename
    stem = filename_stem.replace("_", " ").replace("-", " ")
    return " ".join(w.capitalize() for w in stem.split())


# ── Text Chunking ──────────────────────────────────────────────────────────────

def _chunk_text(
    text: str,
    chunk_size: int = CHUNK_WORD_SIZE,
    overlap:    int = CHUNK_OVERLAP,
) -> list[str]:
    """
    Split text into overlapping fixed-size word chunks.
    Short texts (≤ chunk_size words) return a single chunk.
    """
    words = text.split()
    if not words:
        return []
    if len(words) <= chunk_size:
        return [" ".join(words)]

    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunks.append(" ".join(words[start:end]))
        if end >= len(words):
            break
        start = end - overlap
    return chunks


# ── PDF Parsing ────────────────────────────────────────────────────────────────

def _parse_pdf_text_layer(path: Path) -> str:
    """Extract text using pdfplumber (reads the PDF text layer)."""
    import pdfplumber  # type: ignore

    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def _parse_pdf_ocr(path: Path) -> str:
    """
    Convert PDF pages to images and run Tesseract OCR.
    Only called when the text layer is absent or too short.
    Requires: tesseract-ocr (system), pytesseract, pdf2image, Pillow.
    """
    from pdf2image.info import pdfinfo_from_path # type: ignore
    
    logger.info("Running OCR on %s (scanned PDF detected)", path.name)
    try:
        info = pdfinfo_from_path(str(path))
        num_pages = int(info.get("Pages", 1))
    except Exception:
        num_pages = 1
        
    texts: list[str] = []
    for i in range(1, num_pages + 1):
        pages = convert_from_path(str(path), dpi=200, first_page=i, last_page=i)
        if not pages:
            continue
        page_img = pages[0]
        page_text = pytesseract.image_to_string(page_img, lang="eng")
        if page_text.strip():
            texts.append(page_text)
        logger.debug("OCR page %d/%d: %d chars", i, num_pages, len(page_text))
    return "\n".join(texts)


def _parse_pdf(path: Path) -> tuple[str, bool]:
    """
    Parse a PDF file. Returns (text, ocr_used).

    Strategy:
      1. Try pdfplumber (reads text layer, fast, CPU-only)
      2. If text is shorter than _MIN_TEXT_LENGTH AND OCR_ENABLED:
         fall back to Tesseract OCR
      3. Return whatever text we have (may be empty if all strategies fail)
    """
    text = _parse_pdf_text_layer(path)

    if len(text.strip()) >= _MIN_TEXT_LENGTH:
        return text, False   # Good text layer, no OCR needed

    if not OCR_ENABLED:
        if len(text.strip()) < _MIN_TEXT_LENGTH:
            logger.warning(
                "Very little text extracted from %s (%d chars). "
                "This may be a scanned PDF. Set ENABLE_OCR=true to process it.",
                path.name,
                len(text.strip()),
            )
        return text, False

    # Text layer insufficient — attempt OCR
    try:
        ocr_text = _parse_pdf_ocr(path)
        if len(ocr_text.strip()) > len(text.strip()):
            logger.info(
                "OCR produced more text than text layer for %s (%d vs %d chars)",
                path.name,
                len(ocr_text.strip()),
                len(text.strip()),
            )
            return ocr_text, True
        return text, False
    except Exception as e:
        logger.warning("OCR failed for %s: %s — using text layer fallback", path.name, e)
        return text, False


# ── DOCX Parsing ──────────────────────────────────────────────────────────────

def _parse_docx(path: Path) -> str:
    """Extract plain text from a DOCX file (paragraphs and tables)."""
    from docx import Document  # type: ignore

    doc = Document(str(path))
    texts = []
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text)
    return "\n".join(texts)


# ── Public API ─────────────────────────────────────────────────────────────────

def parse_file(path: Path, candidate_id: str) -> Optional[ParsedCV]:
    """
    Parse a CV file into a ParsedCV object.

    Returns:
        ParsedCV  — if file was parsed and yielded usable text.
        None      — if file type is unsupported or no text could be extracted.

    Raises:
        Any parsing exception — the caller (indexer/run.py) is responsible
        for catching, logging per-file errors, and continuing.
    """
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        logger.warning("Unsupported file type skipped: %s (%s)", path.name, ext)
        return None

    ocr_used = False
    if ext == ".pdf":
        raw_text, ocr_used = _parse_pdf(path)
    else:  # .docx
        raw_text = _parse_docx(path)

    if not raw_text or len(raw_text.strip()) < _MIN_TEXT_LENGTH:
        logger.warning(
            "Insufficient text from %s (%d chars) — skipping. "
            "If this is a scanned PDF, ensure ENABLE_OCR=true and Tesseract is installed.",
            path.name,
            len(raw_text.strip()) if raw_text else 0,
        )
        return None

    name     = _extract_name(raw_text, path.stem)
    metadata = extract_metadata(raw_text)
    chunks   = _chunk_text(raw_text)

    logger.debug(
        "Parsed '%s' → name='%s', %d words, %d chunks, "
        "exp=%s yrs, loc='%s', skills=%d, ocr=%s",
        path.name, name,
        len(raw_text.split()), len(chunks),
        metadata.experience_years,
        metadata.location or "unknown",
        len(metadata.skills),
        ocr_used,
    )

    return ParsedCV(
        candidate_id=candidate_id,
        name=name,
        cv_path=str(path),
        raw_text=raw_text,
        metadata=metadata,
        chunks=chunks,
        ocr_used=ocr_used,
    )
